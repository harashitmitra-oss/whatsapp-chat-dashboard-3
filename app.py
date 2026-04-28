import streamlit as st
import pandas as pd
import numpy as np
import re
import os
import matplotlib.pyplot as plt
import plotly.express as px
from textblob import TextBlob
from wordcloud import WordCloud
from datetime import datetime
import emoji
from collections import Counter
from fpdf import FPDF
from docx import Document

# ✅ OpenAI (v1.x)
from openai import OpenAI

# ✅ All-countries phone → country detection
import phonenumbers
from phonenumbers import geocoder

# -------------------------------
# App Configuration
# -------------------------------
st.set_page_config(page_title="Tetr Commnunity Analysis", layout="wide")

col1, col2 = st.columns([5, 1])
with col1:
    st.title("📊 Tetr Commnunity Analysis")
with col2:
    if os.path.exists("logo.png"):
        st.image("logo.png", width=120)

st.markdown(
    """
<style>
body { background-color: #ffffff; }
[data-testid="stSidebar"] { background-color: #0b3d2e; color: white; }
h1, h2, h3, h4, h5, h6 { color: #0b3d2e; }
.persona-pill {
    display: inline-block;
    padding: 4px 10px;
    border-radius: 999px;
    background: #fff3cd;
    color: #8a5a00;
    font-weight: 700;
    border: 1px solid #ffd966;
}
.student-pill {
    display: inline-block;
    padding: 4px 10px;
    border-radius: 999px;
    background: #e8f5e9;
    color: #0b3d2e;
    font-weight: 700;
    border: 1px solid #a5d6a7;
}
</style>
""",
    unsafe_allow_html=True,
)

# -------------------------------
# Sidebar: Settings
# -------------------------------
st.sidebar.title("⚙️ Settings")
OPENAI_API_KEY = st.sidebar.text_input("🔐 Enter OpenAI API Key", type="password").strip()


@st.cache_resource(show_spinner=False)
def get_ai_client(api_key: str):
    if not api_key:
        return None
    return OpenAI(api_key=api_key)


ai_client = get_ai_client(OPENAI_API_KEY)


def generate_ai_summary(prompt: str):
    if not ai_client:
        return "⚠️ Please enter a valid OpenAI API key in the sidebar."
    try:
        response = ai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": "You are a business intelligence analyst summarizing WhatsApp group discussions.",
                },
                {"role": "user", "content": prompt[:12000]},
            ],
            temperature=0.3,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"AI summary unavailable: {type(e).__name__}: {e}"


# -------------------------------
# Utility Functions
# -------------------------------
def clean_phone_number(num):
    """
    Keeps matching accurate for Indian-style and international WhatsApp numbers.
    Returns the last 10 digits because WhatsApp exports often include country code,
    spaces, brackets, or hyphens while sheet data may not.
    """
    digits = re.sub(r"\D", "", str(num))
    if not digits or digits.lower() == "nan":
        return ""
    return digits[-10:] if len(digits) >= 10 else digits


def safe_text(value):
    if pd.isna(value):
        return ""
    value = str(value).strip()
    if value.lower() in ["nan", "none", "nat"]:
        return ""
    return value


def analyze_sentiment(text):
    polarity = TextBlob(str(text)).sentiment.polarity
    if polarity > 0.1:
        return "Positive"
    elif polarity < -0.1:
        return "Negative"
    else:
        return "Neutral"


def extract_emojis(text):
    return [char for char in str(text) if char in emoji.EMOJI_DATA]


def detect_country(sender: str):
    s = str(sender).strip()
    if not re.search(r"\+?\d", s):
        return "Unknown"
    try:
        if not s.startswith("+"):
            return "Unknown"
        num = phonenumbers.parse(s, None)
        if not phonenumbers.is_valid_number(num):
            return "Unknown"
        country = geocoder.description_for_number(num, "en")
        return country if country else "Unknown"
    except Exception:
        return "Unknown"


def assign_engagement_level_by_percentile(series, high_q=0.80, low_q=0.20):
    if series.empty:
        return series
    hi = series.quantile(high_q)
    lo = series.quantile(low_q)

    def label(x):
        if x >= hi:
            return "High"
        elif x <= lo:
            return "Low"
        return "Medium"

    return series.apply(label)


def first_existing_col(df, possible_names):
    """Find a column by normalized name, making sheet headers tolerant to spaces/case."""
    normalized = {str(c).strip().lower(): c for c in df.columns}
    for name in possible_names:
        key = str(name).strip().lower()
        if key in normalized:
            return normalized[key]
    return None


# -------------------------------
# Student + Persona Directory Loader
# -------------------------------
STUDENT_PHONE_FILE = "Students_phone.xlsx"


@st.cache_data(show_spinner=False)
def load_students_and_personas(path=STUDENT_PHONE_FILE):
    """
    Expected workbook in GitHub repo root:
    Sheet 1: student data with Name, Phone, UG/PG, Batch
    Sheet 2: persona data with Name, Persona Name, Phone, Email, Email 2 (if exists), UG/PG
    """
    empty_students = pd.DataFrame(columns=["Name", "Phone", "UG/PG", "Batch", "PhoneKey"])
    empty_personas = pd.DataFrame(
        columns=[
            "Name",
            "Persona Name",
            "Phone",
            "Email",
            "Email 2 (if exists)",
            "UG/PG",
            "PhoneKey",
        ]
    )

    if not os.path.exists(path):
        return (
            empty_students,
            empty_personas,
            f"⚠️ `{path}` not found in repo root. Add it to GitHub main folder to enable student/persona matching.",
        )

    try:
        sheets = pd.read_excel(path, sheet_name=None)
        sheet_names = list(sheets.keys())
        if not sheet_names:
            return empty_students, empty_personas, f"⚠️ `{path}` has no readable sheets."

        students_raw = sheets[sheet_names[0]].copy()
        personas_raw = sheets[sheet_names[1]].copy() if len(sheet_names) > 1 else pd.DataFrame()

        # Student sheet
        s_name_col = first_existing_col(students_raw, ["Name", "Student Name", "Full Name"])
        s_phone_col = first_existing_col(
            students_raw,
            ["Phone", "Mobile", "Phone Number", "Contact", "Whatsapp", "WhatsApp Number"],
        )
        s_program_col = first_existing_col(students_raw, ["UG/PG", "UG PG", "Program", "Course"])
        s_batch_col = first_existing_col(students_raw, ["Batch", "Batch Name"])

        if s_name_col and s_phone_col:
            students = pd.DataFrame()
            students["Name"] = students_raw[s_name_col].apply(safe_text)
            students["Phone"] = students_raw[s_phone_col]
            students["UG/PG"] = students_raw[s_program_col].apply(safe_text) if s_program_col else ""
            students["Batch"] = students_raw[s_batch_col].apply(safe_text) if s_batch_col else ""
            students["PhoneKey"] = students["Phone"].apply(clean_phone_number)
            students = students[(students["PhoneKey"] != "") & (students["Name"] != "")]
            students = students.drop_duplicates(subset=["PhoneKey"], keep="first")
        else:
            students = empty_students.copy()

        # Persona sheet
        if not personas_raw.empty:
            p_name_col = first_existing_col(personas_raw, ["Name", "Student Name", "Full Name"])
            p_persona_col = first_existing_col(
                personas_raw, ["Persona Name", "Persona", "Alias", "Persona/Name"]
            )
            p_phone_col = first_existing_col(
                personas_raw,
                ["Phone", "Mobile", "Phone Number", "Contact", "Whatsapp", "WhatsApp Number"],
            )
            p_email_col = first_existing_col(personas_raw, ["Email", "Email 1", "Primary Email"])
            p_email2_col = first_existing_col(
                personas_raw,
                ["Email 2 (if exists)", "Email 2", "Secondary Email", "Alternate Email"],
            )
            p_program_col = first_existing_col(personas_raw, ["UG/PG", "UG PG", "Program", "Course"])

            if p_phone_col and (p_persona_col or p_name_col):
                personas = pd.DataFrame()
                personas["Name"] = personas_raw[p_name_col].apply(safe_text) if p_name_col else ""
                personas["Persona Name"] = (
                    personas_raw[p_persona_col].apply(safe_text) if p_persona_col else personas["Name"]
                )
                personas["Phone"] = personas_raw[p_phone_col]
                personas["Email"] = personas_raw[p_email_col].apply(safe_text) if p_email_col else ""
                personas["Email 2 (if exists)"] = (
                    personas_raw[p_email2_col].apply(safe_text) if p_email2_col else ""
                )
                personas["UG/PG"] = personas_raw[p_program_col].apply(safe_text) if p_program_col else ""
                personas["PhoneKey"] = personas["Phone"].apply(clean_phone_number)
                personas = personas[(personas["PhoneKey"] != "") & (personas["Persona Name"] != "")]
                personas = personas.drop_duplicates(subset=["PhoneKey"], keep="first")
            else:
                personas = empty_personas.copy()
        else:
            personas = empty_personas.copy()

        status = f"✅ Loaded `{path}`: {len(students)} student phones and {len(personas)} persona phones."
        return students, personas, status

    except Exception as e:
        return empty_students, empty_personas, f"⚠️ Could not read `{path}`: {type(e).__name__}: {e}"


students_df, personas_df, directory_status = load_students_and_personas()
st.sidebar.caption(directory_status)

# Build maps once
student_phone_map = students_df.set_index("PhoneKey").to_dict("index") if not students_df.empty else {}
persona_phone_map = personas_df.set_index("PhoneKey").to_dict("index") if not personas_df.empty else {}


def enrich_sender(sender):
    phone_key = clean_phone_number(sender)
    persona_row = persona_phone_map.get(phone_key)
    student_row = student_phone_map.get(phone_key)

    is_persona = persona_row is not None
    persona_name = safe_text(persona_row.get("Persona Name", "")) if persona_row else ""
    student_name = ""

    if persona_row and safe_text(persona_row.get("Name", "")):
        student_name = safe_text(persona_row.get("Name", ""))
    elif student_row:
        student_name = safe_text(student_row.get("Name", ""))

    display_name = persona_name if is_persona and persona_name else (student_name if student_name else sender)

    ug_pg = ""
    batch = ""
    email = ""
    email2 = ""

    if student_row:
        ug_pg = safe_text(student_row.get("UG/PG", ""))
        batch = safe_text(student_row.get("Batch", ""))

    if persona_row:
        ug_pg = safe_text(persona_row.get("UG/PG", ug_pg)) or ug_pg
        email = safe_text(persona_row.get("Email", ""))
        email2 = safe_text(persona_row.get("Email 2 (if exists)", ""))

    return pd.Series(
        {
            "PhoneKey": phone_key,
            "DisplayName": display_name,
            "StudentName": student_name,
            "PersonaName": persona_name,
            "IsPersona": is_persona,
            "UG/PG": ug_pg,
            "Batch": batch,
            "Email": email,
            "Email 2": email2,
            "MatchType": "Persona" if is_persona else ("Student" if student_row else "Unmatched"),
        }
    )


# -------------------------------
# WhatsApp Chat Parser (Android + iPhone Robust)
# -------------------------------
def strip_whatsapp_invisible_chars(text):
    """Remove hidden marks that iPhone WhatsApp exports often include."""
    return re.sub(r"[\u200e\u200f\u202a-\u202e\ufeff]", "", str(text)).strip()


def parse_whatsapp_chat(file):
    content = file.read()

    try:
        content = content.decode("utf-8-sig")
    except Exception:
        try:
            content = content.decode("utf-16")
        except Exception:
            content = content.decode("latin-1", errors="ignore")

    lines = content.splitlines()
    data = []

    # Supported formats:
    # Android: 16/03/26, 00:40 - +91 99999 99999: Message
    # Android 12-hour: 16/03/2026, 12:40 AM - Name: Message
    # iPhone: [16/03/26, 00:40:12] Name: Message
    # iPhone 12-hour: [16/03/2026, 12:40:12 AM] Name: Message
    # iPhone without seconds: [16/03/26, 00:40] Name: Message
    patterns = [
        re.compile(
            r"^\[?(\d{1,2}[/-]\d{1,2}[/-]\d{2,4}),\s+"
            r"(\d{1,2}:\d{2}(?::\d{2})?)\s*(AM|PM|am|pm)?\]?"
            r"\s+-\s+(.*?):\s*(.*)$"
        ),
        re.compile(
            r"^\[(\d{1,2}[/-]\d{1,2}[/-]\d{2,4}),\s+"
            r"(\d{1,2}:\d{2}(?::\d{2})?)\s*(AM|PM|am|pm)?\]"
            r"\s+(.*?):\s*(.*)$"
        ),
    ]

    current_message = None

    for raw_line in lines:
        line = strip_whatsapp_invisible_chars(raw_line)
        if not line:
            continue

        match = None
        for pattern in patterns:
            match = pattern.match(line)
            if match:
                break

        if match:
            date, time, meridiem, sender, message = match.groups()

            timestamp_str = f"{date} {time} {meridiem}" if meridiem else f"{date} {time}"
            timestamp = pd.to_datetime(timestamp_str, dayfirst=True, errors="coerce")

            if pd.isna(timestamp) or not str(sender).strip():
                current_message = None
                continue

            current_message = [
                timestamp,
                strip_whatsapp_invisible_chars(sender),
                strip_whatsapp_invisible_chars(message),
            ]
            data.append(current_message)
        else:
            # Multi-line message continuation
            if current_message:
                current_message[2] += " " + line

    df = pd.DataFrame(data, columns=["DateTime", "Sender", "Message"])
    df = df.dropna(subset=["DateTime"])

    if df.empty:
        return df

    df["Date"] = df["DateTime"].dt.date
    df["Time"] = df["DateTime"].dt.time
    df["Hour"] = df["DateTime"].dt.hour
    df["Week"] = df["DateTime"].dt.to_period("W").astype(str)
    df["Month"] = df["DateTime"].dt.to_period("M").astype(str)

    return df


# -------------------------------
# Report Storage
# -------------------------------
DATA_DIR = "stored_reports"
os.makedirs(DATA_DIR, exist_ok=True)


def save_daily_report(group_name, df_summary):
    date_str = datetime.now().strftime("%Y-%m-%d")
    safe_group = re.sub(r"[^A-Za-z0-9_-]+", "_", str(group_name)).strip("_") or "Default_Group"
    file_path = os.path.join(DATA_DIR, f"{safe_group}_{date_str}.csv")
    df_summary.to_csv(file_path, index=False)


def load_historical_reports():
    files = os.listdir(DATA_DIR)
    data = []
    for f in files:
        if f.endswith(".csv"):
            df_hist = pd.read_csv(os.path.join(DATA_DIR, f))
            df_hist["SourceFile"] = f
            data.append(df_hist)
    return pd.concat(data, ignore_index=True) if data else pd.DataFrame()


# -------------------------------
# PDF & Word Report Generation
# -------------------------------
def generate_pdf_report(summary_text, metrics_dict, filename="report.pdf"):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=10)
    pdf.multi_cell(0, 6, "WhatsApp Group Intelligence Report\n\n")
    for k, v in metrics_dict.items():
        pdf.multi_cell(0, 6, f"{k}: {v}")
    pdf.multi_cell(0, 6, "\nAI Summary:\n")
    pdf.multi_cell(0, 6, str(summary_text).encode("latin-1", "replace").decode("latin-1"))
    pdf.output(filename)
    return filename


def generate_word_report(summary_text, metrics_dict, filename="report.docx"):
    doc = Document()
    doc.add_heading("WhatsApp Group Intelligence Report", level=1)
    for k, v in metrics_dict.items():
        doc.add_paragraph(f"{k}: {v}")
    doc.add_heading("AI Summary", level=2)
    doc.add_paragraph(str(summary_text))
    doc.save(filename)
    return filename


# -------------------------------
# Sidebar: Group + Window
# -------------------------------
st.sidebar.header("👥 Group Management")
group_name = st.sidebar.text_input("Enter Group Name", value="Default_Group")

# -------------------------------
# Upload Files
# -------------------------------
st.header("📂 Upload WhatsApp Chat File")
chat_file = st.file_uploader("Upload WhatsApp .txt chat file", type=["txt"])

if not chat_file:
    st.info("📥 Please upload a WhatsApp chat file to begin analysis.")
    st.stop()

# -------------------------------
# Parse + Enrich
# -------------------------------
df = parse_whatsapp_chat(chat_file)

if df.empty:
    st.error(
        "No valid WhatsApp messages could be parsed. Please upload an Android or iPhone WhatsApp .txt export with timestamps and sender names/numbers."
    )
    st.stop()

sender_enrichment = df["Sender"].apply(enrich_sender)
df = pd.concat([df, sender_enrichment], axis=1)
df["Country"] = df["Sender"].apply(detect_country)
df["Sentiment"] = df["Message"].apply(analyze_sentiment)
df["Emojis"] = df["Message"].apply(extract_emojis)
df["EmojiCount"] = df["Emojis"].apply(len)
df["WordCount"] = df["Message"].astype(str).str.split().apply(len)
df["IsMedia"] = df["Message"].astype(str).str.contains(
    r"<Media omitted>|image omitted|video omitted|audio omitted|sticker omitted",
    case=False,
    regex=True,
    na=False,
)
df["IsReactionLike"] = df["Message"].astype(str).str.contains(
    r"reacted|reaction|👍|❤️|😂|🤣|😮|😢|🙏",
    case=False,
    regex=True,
    na=False,
)

# -------------------------------
# Analysis Window Selector
# -------------------------------
st.sidebar.header("🕒 Analysis Window")
window_mode = st.sidebar.selectbox("Choose Window", ["All", "Daily", "Weekly", "Monthly"])

window_value = None
if window_mode == "Daily":
    window_value = st.sidebar.selectbox("Pick Date", sorted(df["Date"].dropna().unique()))
elif window_mode == "Weekly":
    window_value = st.sidebar.selectbox("Pick Week", sorted(df["Week"].dropna().unique()))
elif window_mode == "Monthly":
    window_value = st.sidebar.selectbox("Pick Month", sorted(df["Month"].dropna().unique()))

df_window = df.copy()
if window_mode == "Daily":
    df_window = df[df["Date"] == window_value].copy()
elif window_mode == "Weekly":
    df_window = df[df["Week"] == window_value].copy()
elif window_mode == "Monthly":
    df_window = df[df["Month"] == window_value].copy()

if df_window.empty:
    st.warning("No messages found in the selected window.")
    st.stop()

# -------------------------------
# Core Metrics (Window-based)
# -------------------------------
total_messages = len(df_window)
active_members = df_window["DisplayName"].nunique()
total_members_input = st.sidebar.number_input("Total Members in Group", min_value=1, value=int(active_members))
silent_members = max(int(total_members_input) - int(active_members), 0)
activation_rate = round((active_members / total_members_input) * 100, 2)

# -------------------------------
# Participant Aggregation (Percent-based Engagement)
# -------------------------------
total_msgs_window = len(df_window)
total_days_window = df_window["Date"].nunique()
total_weeks_window = df_window["Week"].nunique()
total_months_window = df_window["Month"].nunique()

participant_stats = df_window.groupby("DisplayName").agg(
    MessageCount=("Message", "count"),
    ActiveDays=("Date", "nunique"),
    ActiveWeeks=("Week", "nunique"),
    ActiveMonths=("Month", "nunique"),
    AvgSentiment=("Sentiment", lambda x: x.mode()[0] if not x.mode().empty else "Neutral"),
    PhoneKey=("PhoneKey", "first"),
    Sender=("Sender", "first"),
    StudentName=("StudentName", "first"),
    PersonaName=("PersonaName", "first"),
    IsPersona=("IsPersona", "max"),
    MatchType=("MatchType", "first"),
    UGPG=("UG/PG", "first"),
    Batch=("Batch", "first"),
    Email=("Email", "first"),
    Email2=("Email 2", "first"),
    EmojiCount=("EmojiCount", "sum"),
    MediaCount=("IsMedia", "sum"),
    ReactionLikeCount=("IsReactionLike", "sum"),
    WordCount=("WordCount", "sum"),
).reset_index()

participant_stats["MessageSharePct"] = (participant_stats["MessageCount"] / max(total_msgs_window, 1)) * 100
participant_stats["ActiveDaysPct"] = (participant_stats["ActiveDays"] / max(total_days_window, 1)) * 100

if window_mode in ["All", "Monthly"]:
    participant_stats["ConsistencyPct"] = (participant_stats["ActiveWeeks"] / max(total_weeks_window, 1)) * 100
elif window_mode == "Weekly":
    participant_stats["ConsistencyPct"] = participant_stats["ActiveDaysPct"]
else:
    participant_stats["ConsistencyPct"] = participant_stats["ActiveDaysPct"]

participant_stats["EngagementIndex"] = (
    0.55 * participant_stats["MessageSharePct"]
    + 0.30 * participant_stats["ActiveDaysPct"]
    + 0.15 * participant_stats["ConsistencyPct"]
).round(2)

participant_stats["Sentiment"] = participant_stats["AvgSentiment"]
participant_stats["EngagementLevel"] = assign_engagement_level_by_percentile(participant_stats["EngagementIndex"])


def calculate_lead_score_v2(row):
    score = row["EngagementIndex"]
    if row["Sentiment"] == "Positive":
        score += 15
    elif row["Sentiment"] == "Negative":
        score -= 10
    if row["EngagementLevel"] == "High":
        score += 15
    elif row["EngagementLevel"] == "Medium":
        score += 5
    return round(score, 2)


participant_stats["LeadScore"] = participant_stats.apply(calculate_lead_score_v2, axis=1)
participant_stats["OverallScore"] = (
    participant_stats["LeadScore"]
    + 0.05 * participant_stats["EmojiCount"]
    + 0.10 * participant_stats["ReactionLikeCount"]
    + 0.02 * participant_stats["WordCount"]
).round(2)

top_lead = (
    participant_stats.sort_values("LeadScore", ascending=False).iloc[0]["DisplayName"]
    if len(participant_stats)
    else "N/A"
)

save_daily_report(group_name, participant_stats)

# -------------------------------
# KPI Panel
# -------------------------------
st.subheader("📌 Key Metrics (Selected Window)")

matched_batch_df = participant_stats[
    (participant_stats["MatchType"].isin(["Student", "Persona"]))
    & (participant_stats["UGPG"].astype(str).str.strip() != "")
    & (participant_stats["UGPG"].astype(str).str.lower().str.strip() != "nan")
    & (participant_stats["Batch"].astype(str).str.strip() != "")
    & (participant_stats["Batch"].astype(str).str.lower().str.strip() != "nan")
].copy()

if not matched_batch_df.empty:
    batch_summary_top = (
        matched_batch_df.groupby(["UGPG", "Batch"], dropna=False)
        .agg(MatchedParticipants=("DisplayName", "nunique"), Messages=("MessageCount", "sum"))
        .reset_index()
        .sort_values(["MatchedParticipants", "Messages"], ascending=False)
        .iloc[0]
    )
    detected_batch_label = f"{batch_summary_top['UGPG']} · Batch {batch_summary_top['Batch']}"
    detected_batch_count = int(batch_summary_top["MatchedParticipants"])
else:
    detected_batch_label = "Not detected"
    detected_batch_count = 0

col0, col1, col2, col3, col4, col5 = st.columns(6)
col0.metric("Detected Batch", detected_batch_label, f"{detected_batch_count} matched")
col1.metric("Activation Rate (%)", activation_rate)
col2.metric("Total Messages", total_messages)
col3.metric("Active Members", active_members)
col4.metric("Silent Members", silent_members)
col5.metric("Top Lead", top_lead)

# -------------------------------
# Student + Persona Matching Overview
# -------------------------------
st.subheader("🧾 Student & Persona Matching Overview")
match_summary = participant_stats["MatchType"].value_counts().reset_index()
match_summary.columns = ["Match Type", "Participants"]

col_m1, col_m2, col_m3, col_m4 = st.columns(4)
col_m1.metric("Matched Students", int((participant_stats["MatchType"] == "Student").sum()))
col_m2.metric("Matched Personas", int((participant_stats["MatchType"] == "Persona").sum()))
col_m3.metric("Unmatched Senders", int((participant_stats["MatchType"] == "Unmatched").sum()))
col_m4.metric("Persona Messages", int(df_window[df_window["IsPersona"]]["Message"].count()))

if not match_summary.empty:
    fig_match = px.pie(
        match_summary,
        names="Match Type",
        values="Participants",
        hole=0.45,
        title="Student / Persona / Unmatched Participants",
    )
    st.plotly_chart(fig_match, use_container_width=True)

match_table = participant_stats[
    [
        "DisplayName",
        "MatchType",
        "StudentName",
        "PersonaName",
        "UGPG",
        "Batch",
        "Email",
        "Email2",
        "MessageCount",
        "EngagementIndex",
        "LeadScore",
        "OverallScore",
    ]
].sort_values(["MatchType", "MessageCount"], ascending=[True, False])

match_table = match_table.rename(columns={"UGPG": "UG/PG", "Email2": "Email 2"})


def highlight_persona_rows(row):
    if row.get("MatchType") == "Persona":
        return ["background-color: #fff3cd; color: #5f4300; font-weight: 700" for _ in row]
    if row.get("MatchType") == "Student":
        return ["background-color: #e8f5e9; color: #0b3d2e" for _ in row]
    return ["" for _ in row]


st.dataframe(match_table.style.apply(highlight_persona_rows, axis=1), use_container_width=True)

# -------------------------------
# Persona Activity Section (Window-aware)
# -------------------------------
st.subheader("🎭 Persona Activity Analytics (Selected Window)")
persona_stats = participant_stats[participant_stats["IsPersona"] == True].copy()

if persona_stats.empty:
    st.info(
        "No persona activity found in the selected window. Make sure persona phone numbers are present in `Students_phone.xlsx` sheet 2 and match the WhatsApp sender numbers."
    )
else:
    col_p1, col_p2, col_p3, col_p4, col_p5 = st.columns(5)
    col_p1.metric("Active Personas", int(persona_stats["DisplayName"].nunique()))
    col_p2.metric("Persona Messages", int(persona_stats["MessageCount"].sum()))
    col_p3.metric("Persona Emojis", int(persona_stats["EmojiCount"].sum()))
    col_p4.metric("Persona Media", int(persona_stats["MediaCount"].sum()))
    col_p5.metric("Avg Persona Score", round(float(persona_stats["OverallScore"].mean()), 2))

    persona_activity_table = persona_stats[
        [
            "DisplayName",
            "StudentName",
            "PersonaName",
            "UGPG",
            "MessageCount",
            "MessageSharePct",
            "ActiveDays",
            "EmojiCount",
            "ReactionLikeCount",
            "MediaCount",
            "Sentiment",
            "EngagementIndex",
            "LeadScore",
            "OverallScore",
        ]
    ].rename(columns={"UGPG": "UG/PG"}).sort_values("OverallScore", ascending=False)

    st.dataframe(
        persona_activity_table.style.apply(
            lambda row: ["background-color: #fff3cd; color: #5f4300; font-weight: 700" for _ in row],
            axis=1,
        ),
        use_container_width=True,
    )

    persona_msg_chart = px.bar(
        persona_activity_table,
        x="DisplayName",
        y="MessageCount",
        text="MessageCount",
        title="Persona Message Count",
    )
    st.plotly_chart(persona_msg_chart, use_container_width=True)

    persona_score_chart = px.bar(
        persona_activity_table,
        x="DisplayName",
        y="OverallScore",
        text="OverallScore",
        title="Persona Overall Score",
    )
    st.plotly_chart(persona_score_chart, use_container_width=True)

    persona_df_window = df_window[df_window["IsPersona"] == True].copy()

    st.markdown("### 😄 Persona Emoji / Reaction Activity")
    persona_emojis = [e for sublist in persona_df_window["Emojis"] for e in sublist]
    persona_emoji_counts = Counter(persona_emojis)
    persona_emoji_df = pd.DataFrame(persona_emoji_counts.most_common(20), columns=["Emoji", "Count"])
    if not persona_emoji_df.empty:
        persona_emoji_df["Pct"] = (persona_emoji_df["Count"] / persona_emoji_df["Count"].sum()) * 100
        st.plotly_chart(
            px.bar(persona_emoji_df, x="Emoji", y="Count", text="Count", title="Top Persona Emojis"),
            use_container_width=True,
        )
    else:
        st.info("No persona emojis found in this selected window.")

    st.markdown("### 📌 Latest Persona Messages")
    latest_persona_cols = [
        "DateTime",
        "DisplayName",
        "StudentName",
        "PersonaName",
        "Message",
        "Sentiment",
        "EmojiCount",
        "IsMedia",
        "IsReactionLike",
    ]
    latest_persona_cols = [c for c in latest_persona_cols if c in persona_df_window.columns]
    latest_persona_msgs = persona_df_window.sort_values("DateTime", ascending=False)[
        latest_persona_cols
    ].head(100)
    latest_persona_msgs = latest_persona_msgs.rename(columns={"IsReactionLike": "ReactionLike"})
    st.dataframe(latest_persona_msgs, use_container_width=True)

    st.markdown("### 📈 Persona Activity Over Time")
    persona_timeline_option = st.radio(
        "Persona timeline view by:",
        ["Daily", "Weekly", "Monthly"],
        horizontal=True,
        key="persona_timeline_option",
    )
    if persona_timeline_option == "Daily":
        persona_timeline = persona_df_window.groupby(["Date", "DisplayName"]).size().reset_index(name="Messages")
        fig_persona_timeline = px.line(
            persona_timeline,
            x="Date",
            y="Messages",
            color="DisplayName",
            title="Persona Daily Activity",
        )
    elif persona_timeline_option == "Weekly":
        persona_timeline = persona_df_window.groupby(["Week", "DisplayName"]).size().reset_index(name="Messages")
        fig_persona_timeline = px.line(
            persona_timeline,
            x="Week",
            y="Messages",
            color="DisplayName",
            title="Persona Weekly Activity",
        )
    else:
        persona_timeline = persona_df_window.groupby(["Month", "DisplayName"]).size().reset_index(name="Messages")
        fig_persona_timeline = px.line(
            persona_timeline,
            x="Month",
            y="Messages",
            color="DisplayName",
            title="Persona Monthly Activity",
        )
    st.plotly_chart(fig_persona_timeline, use_container_width=True)

# -------------------------------
# Time-of-Day Chart (Window-aware)
# -------------------------------
st.subheader("⏰ Time-of-Day Message Activity (Selected Window)")
hourly_counts = df_window.groupby("Hour").size().reset_index(name="Messages")
all_hours = pd.DataFrame({"Hour": list(range(24))})
hourly_counts = all_hours.merge(hourly_counts, on="Hour", how="left").fillna({"Messages": 0})
hourly_counts["HourLabel"] = hourly_counts["Hour"].apply(lambda h: f"{int(h):02d}:00")

peak_hour_row = hourly_counts.sort_values("Messages", ascending=False).iloc[0]
st.caption(
    f"Peak message time in selected filter: {peak_hour_row['HourLabel']} with {int(peak_hour_row['Messages'])} messages."
)
fig_hourly = px.bar(
    hourly_counts,
    x="HourLabel",
    y="Messages",
    text="Messages",
    title="Messages by Hour of Day",
)
st.plotly_chart(fig_hourly, use_container_width=True)

# Optional split by persona/student/unmatched for the same selected filter
hourly_by_type = df_window.groupby(["Hour", "MatchType"]).size().reset_index(name="Messages")
hourly_by_type["HourLabel"] = hourly_by_type["Hour"].apply(lambda h: f"{int(h):02d}:00")
if not hourly_by_type.empty:
    fig_hourly_type = px.bar(
        hourly_by_type,
        x="HourLabel",
        y="Messages",
        color="MatchType",
        barmode="group",
        title="Messages by Hour and Sender Type",
    )
    st.plotly_chart(fig_hourly_type, use_container_width=True)

# -------------------------------
# Engagement Pies
# -------------------------------
st.subheader("📊 Engagement Contribution (Percent-based)")

eng_msg = participant_stats.groupby("EngagementLevel")["MessageSharePct"].sum().reset_index()
fig_eng_share = px.pie(
    eng_msg,
    names="EngagementLevel",
    values="MessageSharePct",
    hole=0.45,
    title="Message Share by Engagement Level (%)",
)
fig_eng_share.update_traces(textinfo="percent+label")
st.plotly_chart(fig_eng_share, use_container_width=True)

eng_index = participant_stats.groupby("EngagementLevel")["EngagementIndex"].sum().reset_index()
eng_index["IndexSharePct"] = (eng_index["EngagementIndex"] / max(eng_index["EngagementIndex"].sum(), 1)) * 100

fig_eng_index = px.pie(
    eng_index,
    names="EngagementLevel",
    values="IndexSharePct",
    hole=0.45,
    title="Engagement Index Share by Level (%)",
)
fig_eng_index.update_traces(textinfo="percent+label")
st.plotly_chart(fig_eng_index, use_container_width=True)

# -------------------------------
# Users by Engagement Level + AI
# -------------------------------
st.subheader("👥 Users by Engagement Level")

for level in ["High", "Medium", "Low"]:
    st.markdown(f"### {level} Engagement Users")
    subset = participant_stats[participant_stats["EngagementLevel"] == level].sort_values(
        "EngagementIndex", ascending=False
    )
    display_cols = [
        "DisplayName",
        "MatchType",
        "StudentName",
        "PersonaName",
        "MessageCount",
        "MessageSharePct",
        "ActiveDaysPct",
        "EngagementIndex",
        "Sentiment",
        "LeadScore",
        "OverallScore",
    ]
    st.dataframe(subset[display_cols].style.apply(highlight_persona_rows, axis=1), use_container_width=True)

    subset_df = df_window[df_window["DisplayName"].isin(subset["DisplayName"])]
    if len(subset_df) > 0:
        sample_size = min(40, len(subset_df))
        sample_msgs = subset_df.sort_values("DateTime").tail(sample_size)["Message"].tolist()
        msgs_text = "\n- " + "\n- ".join([str(m) for m in sample_msgs if str(m).strip()][:40])
    else:
        msgs_text = ""

    ai_prompt = f"""
Analyze the following WhatsApp messages from {level} engagement users.
Summarize key themes, concerns, interests, and any business signals.
Write business-friendly bullet points.

Messages:{msgs_text}
"""
    st.info(f"🤖 AI Insight on {level} Engagement Users:\n\n{generate_ai_summary(ai_prompt)}")

# -------------------------------
# Sentiment Distribution (Window)
# -------------------------------
st.subheader("💬 Sentiment Distribution (Selected Window)")
sent_counts = df_window["Sentiment"].value_counts().reset_index()
sent_counts.columns = ["Sentiment", "Count"]
sent_counts["Pct"] = (sent_counts["Count"] / sent_counts["Count"].sum()) * 100

fig_sent = px.bar(
    sent_counts,
    x="Sentiment",
    y="Pct",
    title="Sentiment Distribution (%)",
    text=sent_counts["Pct"].round(1),
)
st.plotly_chart(fig_sent, use_container_width=True)

# -------------------------------
# Emoji Analysis (Window)
# -------------------------------
st.subheader("😄 Emoji Usage Analysis (Selected Window)")
all_emojis = [e for sublist in df_window["Emojis"] for e in sublist]
emoji_counts = Counter(all_emojis)

emoji_df = pd.DataFrame(emoji_counts.most_common(15), columns=["Emoji", "Count"])
if not emoji_df.empty:
    emoji_df["Pct"] = (emoji_df["Count"] / emoji_df["Count"].sum()) * 100
    fig_emoji = px.bar(
        emoji_df,
        x="Emoji",
        y="Pct",
        title="Top Emojis Used (%)",
        text=emoji_df["Pct"].round(1),
    )
    st.plotly_chart(fig_emoji, use_container_width=True)
else:
    st.info("No emojis found in the selected window.")

st.subheader("📈 Emoji Sentiment Trend Over Time (Window)")
df_exploded = df_window.explode("Emojis")
emoji_sentiment_df = df_exploded.groupby(["Date", "Sentiment"]).size().reset_index(name="Count")
if not emoji_sentiment_df.empty:
    fig_emoji_trend = px.line(
        emoji_sentiment_df,
        x="Date",
        y="Count",
        color="Sentiment",
        title="Emoji-Related Sentiment Trend Over Time",
    )
    st.plotly_chart(fig_emoji_trend, use_container_width=True)
else:
    st.info("Not enough emoji data for trend plotting.")

# -------------------------------
# Topic / Discussion Analysis (Window)
# -------------------------------
st.subheader("🧠 Topic & Discussion Analysis (Selected Window)")
all_text = " ".join(df_window["Message"].astype(str))
if all_text.strip():
    wc = WordCloud(width=900, height=320, background_color="white").generate(all_text)
    fig_wc, ax_wc = plt.subplots(figsize=(12, 4))
    ax_wc.imshow(wc, interpolation="bilinear")
    ax_wc.axis("off")
    st.pyplot(fig_wc)

topic_msgs = df_window["Message"].astype(str).sample(min(250, len(df_window)), random_state=42).tolist()
topic_text = "\n- " + "\n- ".join(topic_msgs[:250])

topic_prompt = f"""
Analyze these WhatsApp group messages and extract:
1) Main topics discussed
2) Key recurring themes
3) Opportunities / risks / concerns
4) Suggested actions for business stakeholders

Messages:{topic_text}
"""
st.info(f"🤖 AI Topic & Theme Analysis:\n\n{generate_ai_summary(topic_prompt)}")

# -------------------------------
# Participant Search & Profile
# -------------------------------
st.subheader("🔍 Participant-Level Analytics")

# ✅ choose from FULL chat users, and timeline uses FULL chat
selected_user = st.selectbox("Select a participant", sorted(df["DisplayName"].unique()))

# Window-based stats for KPIs (if user not in window, show zeros)
user_stats_window = participant_stats[participant_stats["DisplayName"] == selected_user]
if not user_stats_window.empty:
    user_stats = user_stats_window.iloc[0]
else:
    user_stats = {
        "MessageCount": 0,
        "MessageSharePct": 0.0,
        "EngagementIndex": 0.0,
        "EngagementLevel": "N/A",
        "LeadScore": 0.0,
        "OverallScore": 0.0,
        "MatchType": "N/A",
        "StudentName": "",
        "PersonaName": "",
    }

# ✅ FULL chat data for timeline, irrespective of window
user_df = df[df["DisplayName"] == selected_user]

colu1, colu2, colu3, colu4, colu5, colu6 = st.columns(6)
colu1.metric("Message Count (Window)", int(user_stats["MessageCount"]))
colu2.metric("Message Share % (Window)", round(float(user_stats["MessageSharePct"]), 2))
colu3.metric("Engagement Index (Window)", round(float(user_stats["EngagementIndex"]), 2))
colu4.metric("Engagement Level (Window)", user_stats["EngagementLevel"])
colu5.metric("Lead Score (Window)", round(float(user_stats["LeadScore"]), 2))
colu6.metric("Overall Score (Window)", round(float(user_stats.get("OverallScore", 0)), 2))

if str(user_stats.get("MatchType", "")) == "Persona":
    st.markdown(
        f"<span class='persona-pill'>Persona: {user_stats.get('PersonaName', selected_user)}</span>",
        unsafe_allow_html=True,
    )
elif str(user_stats.get("MatchType", "")) == "Student":
    st.markdown(
        f"<span class='student-pill'>Student: {user_stats.get('StudentName', selected_user)}</span>",
        unsafe_allow_html=True,
    )

st.subheader("📅 User Engagement Timeline (Full Chat)")
timeline_option = st.radio("View by:", ["Daily", "Weekly", "Monthly"], horizontal=True)

if timeline_option == "Daily":
    timeline_data = user_df.groupby("Date").size().reset_index(name="Messages")
    fig_user_timeline = px.line(
        timeline_data,
        x="Date",
        y="Messages",
        title=f"{selected_user} - Daily Engagement (Full Chat)",
    )
elif timeline_option == "Weekly":
    timeline_data = user_df.groupby("Week").size().reset_index(name="Messages")
    fig_user_timeline = px.line(
        timeline_data,
        x="Week",
        y="Messages",
        title=f"{selected_user} - Weekly Engagement (Full Chat)",
    )
else:
    timeline_data = user_df.groupby("Month").size().reset_index(name="Messages")
    fig_user_timeline = px.line(
        timeline_data,
        x="Month",
        y="Messages",
        title=f"{selected_user} - Monthly Engagement (Full Chat)",
    )

st.plotly_chart(fig_user_timeline, use_container_width=True)

# -------------------------------
# Overall Engagement Trends (Full Chat)
# -------------------------------
st.subheader("📈 Overall Engagement Trends (Full Chat)")
daily_trend = df.groupby("Date").size().reset_index(name="Messages")
weekly_trend = df.groupby("Week").size().reset_index(name="Messages")
monthly_trend = df.groupby("Month").size().reset_index(name="Messages")

col_t1, col_t2, col_t3 = st.columns(3)
with col_t1:
    st.plotly_chart(px.line(daily_trend, x="Date", y="Messages", title="Daily Engagement (Full Chat)"), use_container_width=True)
with col_t2:
    st.plotly_chart(px.line(weekly_trend, x="Week", y="Messages", title="Weekly Engagement (Full Chat)"), use_container_width=True)
with col_t3:
    st.plotly_chart(px.line(monthly_trend, x="Month", y="Messages", title="Monthly Engagement (Full Chat)"), use_container_width=True)

# -------------------------------
# Heatmap (Window)
# -------------------------------
st.subheader("🔥 Activity Heatmap (Day vs Hour) (Selected Window)")
df_window["DayName"] = df_window["DateTime"].dt.day_name()
day_order = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
heatmap_data = df_window.pivot_table(index="DayName", columns="Hour", values="Message", aggfunc="count").fillna(0)
heatmap_data = heatmap_data.reindex(day_order).dropna(how="all")

fig, ax = plt.subplots(figsize=(12, 4))
im = ax.imshow(heatmap_data.values, aspect="auto", cmap="Blues")
ax.set_xticks(range(len(heatmap_data.columns)))
ax.set_xticklabels(list(heatmap_data.columns))
ax.set_yticks(range(len(heatmap_data.index)))
ax.set_yticklabels(list(heatmap_data.index))
plt.colorbar(im)
st.pyplot(fig)

# -------------------------------
# Country-Level Engagement (Window)
# -------------------------------
st.subheader("🌍 Country-Level Engagement (Selected Window)")
country_counts = df_window["Country"].value_counts().reset_index()
country_counts.columns = ["Country", "Messages"]

hide_unknown = st.checkbox("Hide 'Unknown' country", value=False)
if hide_unknown:
    country_counts = country_counts[country_counts["Country"] != "Unknown"]

if not country_counts.empty:
    fig_country = px.bar(country_counts, x="Country", y="Messages", title="Messages by Country (Window)")
    st.plotly_chart(fig_country, use_container_width=True)
else:
    st.info("No country data available (numbers may not be in international + format).")

# -------------------------------
# Group Comparison Dashboard (Historical)
# -------------------------------
st.subheader("🆚 Group Comparison Dashboard (Historical)")
historical_df = load_historical_reports()
if not historical_df.empty:
    hist_summary = historical_df.groupby("SourceFile").agg(
        AvgMessages=("MessageCount", "mean"),
        AvgLeadScore=("LeadScore", "mean"),
        Users=("DisplayName", "nunique"),
    ).reset_index()

    fig_group_compare = px.bar(
        hist_summary,
        x="SourceFile",
        y="AvgMessages",
        color="AvgLeadScore",
        title="Group-wise Avg Messages vs Avg Lead Score",
    )
    st.plotly_chart(fig_group_compare, use_container_width=True)
else:
    st.info("No historical group data available yet.")

# -------------------------------
# AI Summary (Window)
# -------------------------------
st.subheader("🤖 AI Summary (Selected Window)")
sample_msgs = df_window.sort_values("DateTime").tail(min(350, len(df_window)))["Message"].astype(str).tolist()
sample_text = "\n- " + "\n- ".join(sample_msgs)

ai_prompt_daily = f"""
Summarize the selected WhatsApp conversation window for business stakeholders.
Include:
- Overall tone
- Key topics
- Engagement quality
- Opportunities / risks
- Suggested next actions

Messages:{sample_text}
"""
ai_daily_summary = generate_ai_summary(ai_prompt_daily)
st.success(ai_daily_summary)

st.subheader("🧠 AI Insights & High Intent Signals (Selected Window)")
ai_prompt_insights = f"""
Analyze these WhatsApp messages and identify:
1) Emerging trends
2) Issues/complaints
3) Opportunities for conversion
4) Users showing high buying intent (and why)

Messages:{sample_text}
"""
st.info(generate_ai_summary(ai_prompt_insights))

# -------------------------------
# Automated Report Generation
# -------------------------------
st.subheader("📄 Automated Report Generation")
metrics_dict = {
    "Window": f"{window_mode} {window_value if window_value else ''}".strip(),
    "Activation Rate (%)": activation_rate,
    "Total Messages": total_messages,
    "Active Members": active_members,
    "Silent Members": silent_members,
    "Top Lead": top_lead,
    "Detected Batch": detected_batch_label,
    "Detected Batch Matched Count": detected_batch_count,
    "Persona Messages": int(df_window[df_window["IsPersona"]]["Message"].count()),
}

if st.button("⬇️ Download PDF Report"):
    pdf_file = generate_pdf_report(ai_daily_summary, metrics_dict, filename="whatsapp_report.pdf")
    with open(pdf_file, "rb") as f:
        st.download_button("Download PDF", f, file_name="whatsapp_report.pdf")

if st.button("⬇️ Download Word Report"):
    word_file = generate_word_report(ai_daily_summary, metrics_dict, filename="whatsapp_report.docx")
    with open(word_file, "rb") as f:
        st.download_button("Download Word", f, file_name="whatsapp_report.docx")
